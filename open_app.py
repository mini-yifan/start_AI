from docx import Document
import os
import subprocess
import sys
import time
import pyautogui

def create_and_open_word_doc(file_name=None):
    """
    创建一个新的空白Word文档并尝试打开它。

    Args:
        file_name (str): 要创建的Word文档的名称。
    """
    try:
        if file_name:
            # 1. 创建一个新的空白文档
            doc = Document()
            # 可以选择添加一些初始内容，这里保持完全空白
            # doc.add_paragraph("这是一个新文档。")

            # 2. 保存文档
            # 保存到桌面
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            full_path = os.path.join(desktop_path, file_name)
            # 确保文件名唯一
            counter = 1
            original_path = full_path
            while os.path.exists(full_path):
                name_without_ext = os.path.splitext(original_path)[0]
                ext = os.path.splitext(original_path)[1]
                new_name = f"{os.path.basename(name_without_ext)}_{counter}{ext}"
                full_path = os.path.join(desktop_path, new_name)
                counter += 1
            doc.save(full_path)
            print(f"文档 '{full_path}' 已创建。")

        else:
            # 1. 创建一个新的空白文档
            doc = Document()
            # 可以选择添加一些初始内容，这里保持完全空白
            # doc.add_paragraph("这是一个新文档。")

            # 2. 保存文档
            # 保存到桌面
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            full_path = os.path.join(desktop_path, "new.docx")
            # 确保文件名唯一
            counter = 1
            original_path = full_path
            while os.path.exists(full_path):
                name_without_ext = os.path.splitext(original_path)[0]
                ext = os.path.splitext(original_path)[1]
                new_name = f"{os.path.basename(name_without_ext)}_{counter}{ext}"
                full_path = os.path.join(desktop_path, new_name)
                counter += 1
            doc.save(full_path)
            print(f"文档 '{full_path}' 已创建。")

        # 4. 根据操作系统打开文件
        # 使用 subprocess.Popen 可以更好地控制新启动的进程
        if sys.platform.startswith('darwin'):  # macOS
            # 使用 'open' 命令
            subprocess.Popen(['open', full_path])
            return "创建并打开文档成功。"
        elif os.name == 'nt':  # Windows
            # 使用 'start' 命令，注意 '/WAIT' 会等待程序关闭，这里不需要
            # 'start' 是 cmd 的内置命令，需要通过 'cmd /c' 调用
            # '/b' 参数可以在后台运行，但不创建新窗口，我们不使用它来确保窗口打开
            subprocess.Popen(['cmd', '/c', 'start', '', full_path], shell=True)
            return "创建并打开文档成功。"
        elif os.name == 'posix':  # Linux
            # 尝试使用 'xdg-open' 命令
            subprocess.Popen(['xdg-open', full_path])
            return "创建并打开文档成功。"
        else:
            print(f"不支持的操作系统: {sys.platform}")
            return "创建或打开文档时出错。"

    except Exception as e:
        print(f"创建或打开文档时出错: {e}")
        return "创建或打开文档时出错。"

def open_netease_music():
    """
    在Windows系统上打开网易云音乐客户端
    该函数会尝试查找并启动网易云音乐程序
    """
    try:
        # 直接使用start命令启动
        subprocess.run(["start", "cloudmusic"], shell=True, check=True)
        print("已尝试启动网易云音乐")
        time.sleep(2)
        return True
    except subprocess.CalledProcessError:
        print("无法通过start命令启动网易云音乐")
        return False


# --- 主程序 ---
if __name__ == "__main__":
    # 指定文件名
    #document_name = "new_document.docx"
    # 调用函数创建并打开文档
    return_content = create_and_open_word_doc()
    print(return_content)




